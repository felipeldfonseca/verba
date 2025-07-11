"""
DOCX export module for generating Word documents.

This module provides functionality to generate DOCX files with the structured
content from the meeting summary (Resumo, Decisões, Próximas Ações, Transcrição).
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Union
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn


logger = logging.getLogger(__name__)


class DocxExporter:
    """DOCX document exporter."""
    
    def __init__(self, template_path: Optional[Union[str, Path]] = None):
        """
        Initialize the DOCX exporter.
        
        Args:
            template_path: Optional path to a DOCX template file
        """
        self.template_path = template_path
        
    def create_document(
        self,
        summary_result,
        meeting_title: str = "Ata de Reunião",
        company_name: str = "Verba",
        output_path: Optional[Union[str, Path]] = None
    ) -> str:
        """
        Create a DOCX document from summary result.
        
        Args:
            summary_result: SummaryResult object from GPT summarizer
            meeting_title: Title of the meeting
            company_name: Company name for header
            output_path: Output file path (defaults to auto-generated)
            
        Returns:
            Path to the generated DOCX file
        """
        # Create document
        if self.template_path and Path(self.template_path).exists():
            doc = Document(str(self.template_path))
        else:
            doc = Document()
            
        # Set up document properties
        doc.core_properties.title = meeting_title
        doc.core_properties.author = "Verba - Gerador Automático de Atas"
        doc.core_properties.created = datetime.now()
        
        # Add header
        self._add_header(doc, meeting_title, company_name)
        
        # Add summary sections
        self._add_summary_sections(doc, summary_result)
        
        # Generate output path if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"ata_{timestamp}.docx"
            
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(str(output_path))
        
        logger.info(f"DOCX document saved to {output_path}")
        return str(output_path)
    
    def _add_header(self, doc: Document, meeting_title: str, company_name: str):
        """
        Add header section to the document.
        
        Args:
            doc: Document object
            meeting_title: Title of the meeting
            company_name: Company name
        """
        # Company name
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(company_name)
        company_run.font.size = Pt(16)
        company_run.font.bold = True
        company_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Meeting title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(meeting_title)
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add separator
        doc.add_paragraph("_" * 50).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
    
    def _add_summary_sections(self, doc: Document, summary_result):
        """
        Add the four main sections to the document.
        
        Args:
            doc: Document object
            summary_result: SummaryResult object
        """
        # 1. Resumo Executivo
        self._add_section_heading(doc, "Resumo Executivo")
        self._add_paragraph(doc, summary_result.resumo_executivo)
        doc.add_paragraph()
        
        # 2. Decisões
        self._add_section_heading(doc, "Decisões")
        if summary_result.decisoes:
            for i, decisao in enumerate(summary_result.decisoes, 1):
                self._add_bullet_point(doc, decisao)
        else:
            self._add_paragraph(doc, "*(nenhuma)*", italic=True)
        doc.add_paragraph()
        
        # 3. Próximas Ações
        self._add_section_heading(doc, "Próximas Ações")
        if summary_result.proximas_acoes:
            self._add_actions_table(doc, summary_result.proximas_acoes)
        else:
            self._add_paragraph(doc, "*(nenhuma)*", italic=True)
        doc.add_paragraph()
        
        # 4. Transcrição Completa
        self._add_section_heading(doc, "Transcrição Completa")
        self._add_paragraph(doc, summary_result.transcricao_completa, font_size=10)
        
        # Add processing info
        doc.add_page_break()
        self._add_processing_info(doc, summary_result)
    
    def _add_section_heading(self, doc: Document, heading: str):
        """
        Add a section heading.
        
        Args:
            doc: Document object
            heading: Heading text
        """
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(heading)
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _add_paragraph(
        self, 
        doc: Document, 
        text: str, 
        font_size: int = 11,
        italic: bool = False,
        bold: bool = False
    ):
        """
        Add a paragraph with formatting.
        
        Args:
            doc: Document object
            text: Paragraph text
            font_size: Font size in points
            italic: Whether to make text italic
            bold: Whether to make text bold
        """
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        run.font.italic = italic
        run.font.bold = bold
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _add_bullet_point(self, doc: Document, text: str):
        """
        Add a bullet point.
        
        Args:
            doc: Document object
            text: Bullet point text
        """
        para = doc.add_paragraph()
        run = para.add_run(f"• {text}")
        run.font.size = Pt(11)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _add_actions_table(self, doc: Document, actions: List[Dict[str, str]]):
        """
        Add a table for próximas ações.
        
        Args:
            doc: Document object
            actions: List of action dictionaries
        """
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Responsável'
        header_cells[1].text = 'Ação'
        header_cells[2].text = 'Prazo'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add action rows
        for action in actions:
            row_cells = table.add_row().cells
            row_cells[0].text = action.get('responsavel', '')
            row_cells[1].text = action.get('acao', '')
            row_cells[2].text = action.get('prazo', '')
        
        # Set column widths
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(1.5)
    
    def _add_processing_info(self, doc: Document, summary_result):
        """
        Add processing information section.
        
        Args:
            doc: Document object
            summary_result: SummaryResult object
        """
        self._add_section_heading(doc, "Informações de Processamento")
        
        info_text = f"""
Este documento foi gerado automaticamente pelo sistema Verba.

• Tokens utilizados: {summary_result.tokens_used:,}
• Tempo de processamento: {summary_result.processing_time:.2f} segundos
• Data de geração: {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}

Para dúvidas ou sugestões, entre em contato com a equipe de desenvolvimento.
        """.strip()
        
        self._add_paragraph(doc, info_text, font_size=10, italic=True)


def export_to_docx(
    summary_result,
    meeting_title: str = "Ata de Reunião",
    company_name: str = "Verba",
    output_path: Optional[Union[str, Path]] = None,
    template_path: Optional[Union[str, Path]] = None
) -> str:
    """
    Convenience function to export summary to DOCX.
    
    Args:
        summary_result: SummaryResult object from GPT summarizer
        meeting_title: Title of the meeting
        company_name: Company name for header
        output_path: Output file path (defaults to auto-generated)
        template_path: Optional path to a DOCX template file
        
    Returns:
        Path to the generated DOCX file
    """
    exporter = DocxExporter(template_path)
    return exporter.create_document(
        summary_result, meeting_title, company_name, output_path
    )


def create_template_docx(output_path: Union[str, Path]) -> str:
    """
    Create a basic DOCX template for the meeting minutes.
    
    Args:
        output_path: Path to save the template
        
    Returns:
        Path to the created template
    """
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('Custom Title', 1)  # 1 = PARAGRAPH
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    
    # Heading style
    heading_style = styles.add_style('Custom Heading', 1)
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    
    # Add template content
    doc.add_paragraph("{{COMPANY_NAME}}", style='Custom Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("{{MEETING_TITLE}}", style='Custom Heading').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Data: {{DATE}}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("_" * 50).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # Add placeholders for sections
    doc.add_paragraph("{{RESUMO_EXECUTIVO}}")
    doc.add_paragraph("{{DECISOES}}")
    doc.add_paragraph("{{PROXIMAS_ACOES}}")
    doc.add_paragraph("{{TRANSCRICAO_COMPLETA}}")
    
    # Save template
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    logger.info(f"DOCX template saved to {output_path}")
    return str(output_path)


def format_actions_for_docx(actions: List[Dict[str, str]]) -> str:
    """
    Format actions list for DOCX insertion.
    
    Args:
        actions: List of action dictionaries
        
    Returns:
        Formatted string
    """
    if not actions:
        return "*(nenhuma)*"
    
    formatted = "| Responsável | Ação | Prazo |\n"
    formatted += "|-------------|------|---------|\n"
    
    for action in actions:
        responsavel = action.get('responsavel', '')
        acao = action.get('acao', '')
        prazo = action.get('prazo', '')
        formatted += f"| {responsavel} | {acao} | {prazo} |\n"
    
    return formatted 